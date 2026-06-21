"""PEERLESS.AI - Multi-agent scientific peer-review assistant"""
from __future__ import annotations
import io, re, math, requests
import streamlit as st
from scipy import stats as sp

st.set_page_config(page_title="PEERLESS.AI", page_icon="magnifying_glass", layout="wide")

# ── Secrets ───────────────────────────────────────────────────────────────────
def _secret(key: str, default: str = "") -> str:
    try:
        return st.secrets[key]
    except Exception:
        return default

GROQ_KEY = _secret("GROQ_API_KEY")
CROSSREF_MAILTO = _secret("CROSSREF_MAILTO", "demo@example.com")
GROQ_BASE = "https://api.groq.com/openai/v1"
SMART_MODEL = "llama-3.3-70b-versatile"

# ── CSS ───────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
/* ── Global ── */
html, body, [class*="css"] { font-family: 'Inter', 'Segoe UI', sans-serif; }

/* ── Hero banner ── */
.hero {
    background: linear-gradient(135deg, #0C2340 0%, #1565C0 60%, #1E88E5 100%);
    border-radius: 14px;
    padding: 32px 36px 28px;
    margin-bottom: 24px;
    color: white;
}
.hero h1 {
    font-size: 2.2rem;
    font-weight: 800;
    letter-spacing: -0.5px;
    margin: 0 0 6px 0;
    color: white !important;
}
.hero p {
    font-size: 0.97rem;
    color: #A8C5D8;
    margin: 0;
    line-height: 1.6;
}
.hero .tag {
    display: inline-block;
    background: rgba(168,197,216,0.2);
    border: 1px solid rgba(168,197,216,0.4);
    color: #A8C5D8;
    font-size: 0.78rem;
    font-weight: 600;
    padding: 3px 10px;
    border-radius: 20px;
    margin: 10px 4px 0 0;
}

/* ── Finding cards ── */
.finding-high {
    border-left: 4px solid #DC2626;
    background: #FEF2F2;
    padding: 13px 16px;
    border-radius: 8px;
    margin: 8px 0;
    box-shadow: 0 1px 4px rgba(220,38,38,0.08);
}
.finding-medium {
    border-left: 4px solid #D97706;
    background: #FFFBEB;
    padding: 13px 16px;
    border-radius: 8px;
    margin: 8px 0;
    box-shadow: 0 1px 4px rgba(217,119,6,0.08);
}
.finding-low {
    border-left: 4px solid #1565C0;
    background: #EFF6FF;
    padding: 13px 16px;
    border-radius: 8px;
    margin: 8px 0;
    box-shadow: 0 1px 4px rgba(21,101,192,0.08);
}
.finding-info {
    border-left: 4px solid #9DBFCA;
    background: #F4F8FC;
    padding: 13px 16px;
    border-radius: 8px;
    margin: 8px 0;
    box-shadow: 0 1px 3px rgba(13,31,45,0.05);
}

/* ── Severity badges ── */
.badge {
    padding: 3px 9px;
    border-radius: 5px;
    font-size: 10px;
    font-weight: 800;
    color: white;
    letter-spacing: 0.5px;
    text-transform: uppercase;
}
.badge-high   { background: #DC2626; }
.badge-medium { background: #D97706; }
.badge-low    { background: #1565C0; }
.badge-info   { background: #9DBFCA; color: #0C2340; }

/* ── Metric cards ── */
[data-testid="metric-container"] {
    background: white;
    border: 1px solid #D6E8F4;
    border-radius: 10px;
    padding: 14px 18px !important;
    box-shadow: 0 1px 4px rgba(13,31,45,0.06);
}
[data-testid="metric-container"] label {
    font-size: 0.72rem !important;
    font-weight: 600 !important;
    color: #1565C0 !important;
    text-transform: uppercase;
    letter-spacing: 0.5px;
}
[data-testid="metric-container"] [data-testid="stMetricValue"] {
    font-size: 1.6rem !important;
    font-weight: 800 !important;
    color: #0C2340 !important;
}

/* ── Score bars ── */
.score-wrap { margin: 4px 0 16px; }
.score-row { display: flex; align-items: center; gap: 12px; margin-bottom: 6px; }
.score-label { font-size: 0.78rem; font-weight: 600; color: #0C2340; width: 200px; flex-shrink: 0; }
.score-track {
    flex: 1; height: 8px; background: #D6E8F4; border-radius: 4px; overflow: hidden;
}
.score-fill {
    height: 100%; border-radius: 4px;
    background: linear-gradient(90deg, #1565C0, #1E88E5);
    transition: width 0.4s ease;
}
.score-fill.ok  { background: linear-gradient(90deg, #16a34a, #22c55e); }
.score-fill.warn{ background: linear-gradient(90deg, #d97706, #f59e0b); }
.score-fill.bad { background: linear-gradient(90deg, #dc2626, #f87171); }
.score-num { font-size: 0.78rem; font-weight: 700; color: #1565C0; width: 32px; text-align: right; }

/* ── Big score display ── */
.score-hero {
    display: flex; align-items: baseline; gap: 8px; margin: 0 0 4px;
}
.score-hero .num {
    font-size: 2.4rem; font-weight: 800; color: #0C2340; line-height: 1;
}
.score-hero .denom { font-size: 1.1rem; color: #9DBFCA; font-weight: 600; }
.score-hero .pct {
    font-size: 0.9rem; font-weight: 600; padding: 3px 10px;
    border-radius: 20px; margin-left: 6px;
}
.score-hero .pct.ok   { background: #dcfce7; color: #16a34a; }
.score-hero .pct.warn { background: #fef3c7; color: #d97706; }
.score-hero .pct.bad  { background: #fee2e2; color: #dc2626; }

/* ── Summary verdict card ── */
.verdict {
    border-radius: 10px;
    padding: 20px 24px;
    margin: 16px 0 20px;
    display: flex; gap: 24px; align-items: flex-start;
}
.verdict.high   { background: #f0fdf4; border: 1px solid #bbf7d0; }
.verdict.medium { background: #fffbeb; border: 1px solid #fde68a; }
.verdict.low    { background: #fef2f2; border: 1px solid #fecaca; }
.verdict-label {
    font-size: 0.7rem; font-weight: 700; text-transform: uppercase;
    letter-spacing: 0.8px; margin-bottom: 2px;
}
.verdict.high   .verdict-label { color: #16a34a; }
.verdict.medium .verdict-label { color: #d97706; }
.verdict.low    .verdict-label { color: #dc2626; }
.verdict-title { font-size: 1.1rem; font-weight: 700; color: #0C2340; margin-bottom: 4px; }
.verdict-sub   { font-size: 0.83rem; color: #64748b; }
.verdict-stats { display: flex; gap: 16px; margin-top: 8px; }
.vstat { text-align: center; }
.vstat .vn { font-size: 1.4rem; font-weight: 800; color: #0C2340; line-height: 1; }
.vstat .vl { font-size: 0.68rem; text-transform: uppercase; letter-spacing: 0.5px; color: #9DBFCA; }

/* ── Agent cards (empty state) ── */
.agent-grid {
    display: grid; grid-template-columns: repeat(auto-fill, minmax(220px, 1fr));
    gap: 12px; margin: 20px 0;
}
.agent-card {
    background: white; border: 1px solid #D6E8F4; border-radius: 10px;
    padding: 16px; box-shadow: 0 1px 3px rgba(13,31,45,0.05);
}
.agent-card .ac-num {
    display: inline-flex; align-items: center; justify-content: center;
    width: 24px; height: 24px; border-radius: 50%;
    background: #1565C0; color: white; font-size: 0.72rem; font-weight: 700;
    margin-bottom: 8px;
}
.agent-card .ac-name { font-size: 0.88rem; font-weight: 700; color: #0C2340; margin-bottom: 3px; }
.agent-card .ac-desc { font-size: 0.78rem; color: #64748b; line-height: 1.4; }

/* ── Upload area ── */
.upload-hint {
    text-align: center; padding: 10px 0 6px;
    font-size: 0.85rem; color: #64748b;
}

/* ── Tab section label ── */
.tab-header {
    display: flex; justify-content: space-between; align-items: center;
    padding: 12px 0 8px; border-bottom: 1px solid #D6E8F4; margin-bottom: 14px;
}
.tab-header .th-title { font-size: 0.9rem; font-weight: 700; color: #0C2340; }
.tab-header .th-count {
    font-size: 0.78rem; font-weight: 600; padding: 3px 10px;
    border-radius: 20px; background: #D6E8F4; color: #1565C0;
}

/* ── Disclaimer footer ── */
.disclaimer {
    background: #D6E8F4;
    border-left: 4px solid #1565C0;
    border-radius: 8px;
    padding: 12px 16px;
    font-size: 0.82rem;
    color: #0C2340;
    margin-top: 20px;
}

/* ── Sidebar ── */
section[data-testid="stSidebar"] {
    background: #0C2340 !important;
}
section[data-testid="stSidebar"] * {
    color: #A8C5D8 !important;
}
section[data-testid="stSidebar"] h3,
section[data-testid="stSidebar"] strong {
    color: white !important;
}
section[data-testid="stSidebar"] table {
    font-size: 0.83rem;
}

/* ── Tabs ── */
[data-testid="stTabs"] button {
    font-weight: 600;
    font-size: 0.82rem;
}
[data-testid="stTabs"] button[aria-selected="true"] {
    color: #1565C0 !important;
    border-bottom-color: #1565C0 !important;
}

/* ── Upload zone ── */
[data-testid="stFileUploader"] {
    border: 2px dashed #9DBFCA !important;
    border-radius: 10px;
    background: white;
}

/* ── Progress bar ── */
[data-testid="stProgress"] > div > div {
    background: linear-gradient(90deg, #1565C0, #1E88E5) !important;
    border-radius: 4px;
}

h1 { margin-bottom: 0 !important; }

/* ── Fixed footer ── */
.footer {
    position: fixed;
    bottom: 0;
    left: 0;
    width: 100%;
    background: white;
    border-top: 1px solid #D6E8F4;
    padding: 8px 24px;
    display: flex;
    align-items: center;
    gap: 14px;
    z-index: 999;
}
.footer span {
    font-size: 0.78rem;
    color: #9DBFCA;
}
.footer strong { color: #1565C0; }
.footer a {
    font-size: 0.78rem;
    color: #1565C0;
    text-decoration: none;
    font-weight: 600;
    border: 1px solid #D6E8F4;
    border-radius: 6px;
    padding: 2px 10px;
    background: #F4F8FC;
}
.footer a:hover { background: #D6E8F4; }
/* Push main content up so footer doesn't overlap last element */
.main .block-container { padding-bottom: 60px !important; }
</style>
""", unsafe_allow_html=True)

# ── File text extraction ──────────────────────────────────────────────────────
def _extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        return "\n".join(p.extract_text() or "" for p in reader.pages)
    except Exception as e:
        st.error(f"PDF parse error: {e}")
        return ""

def _extract_docx(data: bytes) -> str:
    try:
        import docx as _docx
        doc = _docx.Document(io.BytesIO(data))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        return "\n".join(parts)
    except Exception as e:
        st.error(f"DOCX parse error: {e}")
        return ""

def extract_text(data: bytes, filename: str = "") -> str:
    if filename.lower().endswith(".docx"):
        return _extract_docx(data)
    return _extract_pdf(data)

# ── GRIM ─────────────────────────────────────────────────────────────────────
def grim_check(mean_str: str, n: int) -> bool:
    try:
        M = float(mean_str)
        d = len(mean_str.split(".")[-1]) if "." in mean_str else 0
        S = round(M * n)
        return round(S / n, d) == M
    except Exception:
        return True

# ── Statcheck ─────────────────────────────────────────────────────────────────
def recompute_p(stat_type: str, value: float, df1: int,
                df2: int | None = None, n: int | None = None) -> float | None:
    try:
        if stat_type == "t":
            return float(2 * sp.t.sf(abs(value), df1))
        if stat_type == "F":
            return float(sp.f.sf(value, df1, df2))
        if stat_type == "chi2":
            return float(sp.chi2.sf(value, df1))
        if stat_type == "r" and n:
            t = value * math.sqrt((n - 2) / max(1e-10, 1 - value ** 2))
            return float(2 * sp.t.sf(abs(t), n - 2))
    except Exception:
        pass
    return None

# ── Regex extractors ──────────────────────────────────────────────────────────
_MEAN_RE  = re.compile(r'M\s*=\s*(\d+\.\d+)[^.]*?[nN]\s*=\s*(\d+)', re.I)
_TTEST_RE = re.compile(r't\s*\(\s*(\d+)\s*\)\s*=\s*(-?\d+\.?\d*)\s*[,;]\s*p\s*([<>=]\s*\.?\d+)', re.I)
_FTEST_RE = re.compile(r'F\s*\(\s*(\d+)\s*,\s*(\d+)\s*\)\s*=\s*(\d+\.?\d*)\s*[,;]\s*p\s*([<>=]\s*\.?\d+)', re.I)
_CHI2_RE  = re.compile(r'[xXχ]2?\s*\(\s*(\d+)\s*\)\s*=\s*(\d+\.?\d*)\s*[,;]\s*p\s*([<>=]\s*\.?\d+)', re.I)
_DOI_RE   = re.compile(r'10\.\d{4,9}/[^\s,;)\]"]+')

def _parse_p(s: str) -> float | None:
    s = re.sub(r'[<>=\s]', '', s)
    if s.startswith("."):
        s = "0" + s
    try:
        return float(s)
    except Exception:
        return None

def extract_dois(text: str) -> list[str]:
    return list(dict.fromkeys(d.rstrip(".,;)") for d in _DOI_RE.findall(text)))

# ── LLM ───────────────────────────────────────────────────────────────────────
def call_groq(prompt: str) -> str:
    if not GROQ_KEY:
        return ""
    try:
        from openai import OpenAI
        client = OpenAI(api_key=GROQ_KEY, base_url=GROQ_BASE)
        r = client.chat.completions.create(
            model=SMART_MODEL,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=1200,
            temperature=0.1,
        )
        return r.choices[0].message.content or ""
    except Exception:
        return ""

# ── Agent 1: Statistical Integrity ───────────────────────────────────────────
def run_statistical_agent(text: str) -> list[dict]:
    findings: list[dict] = []

    for m in _MEAN_RE.finditer(text):
        mean_str, n = m.group(1), int(m.group(2))
        if not grim_check(mean_str, n):
            findings.append(dict(
                severity="high", agent="Statistical Integrity",
                title=f"GRIM inconsistency: M={mean_str}, n={n}",
                detail=(f"No integer sum divided by {n} can round to {mean_str}. "
                        "This mean is arithmetically impossible given the reported sample size."),
                flag=True,
            ))

    stat_claims = []
    for m in _TTEST_RE.finditer(text):
        stat_claims.append(dict(type="t", df=int(m.group(1)), val=float(m.group(2)),
                                p_str=m.group(3), raw=m.group(0)[:70]))
    for m in _FTEST_RE.finditer(text):
        stat_claims.append(dict(type="F", df1=int(m.group(1)), df2=int(m.group(2)),
                                val=float(m.group(3)), p_str=m.group(4), raw=m.group(0)[:70]))
    for m in _CHI2_RE.finditer(text):
        stat_claims.append(dict(type="chi2", df=int(m.group(1)), val=float(m.group(2)),
                                p_str=m.group(3), raw=m.group(0)[:70]))

    for c in stat_claims:
        p_rep = _parse_p(c["p_str"])
        if p_rep is None:
            continue
        if c["type"] == "t":
            p_calc = recompute_p("t", c["val"], c["df"])
        elif c["type"] == "F":
            p_calc = recompute_p("F", c["val"], c["df1"], c["df2"])
        else:
            p_calc = recompute_p("chi2", c["val"], c["df"])
        if p_calc is None:
            continue
        diff = abs(p_calc - p_rep)
        crosses = (p_rep < 0.05) != (p_calc < 0.05)
        if crosses:
            findings.append(dict(
                severity="high", agent="Statistical Integrity",
                title=f"Significance error: {c['raw']}",
                detail=(f"Reported p={p_rep:.4f}, recomputed p={p_calc:.4f}. "
                        "The significance conclusion changes at alpha=0.05."),
                flag=True,
            ))
        elif diff > 0.01:
            findings.append(dict(
                severity="medium", agent="Statistical Integrity",
                title=f"p-value inconsistency: {c['raw']}",
                detail=f"Reported p={p_rep:.4f}, recomputed p={p_calc:.4f} (difference={diff:.4f}).",
                flag=True,
            ))

    if not findings:
        findings.append(dict(
            severity="info", agent="Statistical Integrity",
            title="No statistical inconsistencies detected",
            detail="GRIM check and p-value recomputation found no issues in the extracted statistics.",
            flag=False,
        ))
    return findings

# ── Agent 2: Citation Verifier ────────────────────────────────────────────────
def run_citation_agent(text: str) -> list[dict]:
    dois = extract_dois(text)
    if not dois:
        return [dict(severity="info", agent="Citation Verifier",
                     title="No DOIs found in paper",
                     detail="No DOI patterns (10.xxxx/...) were detected in the text.",
                     flag=False)]
    findings: list[dict] = []
    crossref_available = True
    for doi in dois[:15]:
        try:
            r = requests.get(
                f"https://api.crossref.org/works/{doi}",
                headers={"User-Agent": f"PEERLESS.AI/1.0 (mailto:{CROSSREF_MAILTO})"},
                timeout=8,
            )
            if r.status_code == 200:
                data = r.json().get("message", {})
                title = (data.get("title") or [""])[0]
                retracted = (
                    any(u.get("type") == "retraction" for u in data.get("update-to", []))
                    or "RETRACTED" in title.upper()
                )
                if retracted:
                    findings.append(dict(
                        severity="high", agent="Citation Verifier",
                        title=f"RETRACTED paper cited: {doi}",
                        detail=f"Title: {title[:120]}. This paper has been retracted. Citing retracted work can invalidate conclusions.",
                        flag=True,
                    ))
                else:
                    findings.append(dict(
                        severity="info", agent="Citation Verifier",
                        title=f"DOI verified: {doi}",
                        detail=f"Title: {title[:120] or '(no title)'}",
                        flag=False,
                    ))
            elif r.status_code == 404:
                findings.append(dict(
                    severity="high", agent="Citation Verifier",
                    title=f"DOI not found: {doi}",
                    detail="Crossref returned 404. This reference may be fabricated, mistyped, or unregistered.",
                    flag=True,
                ))
            elif r.status_code == 429:
                crossref_available = False
                findings.append(dict(
                    severity="low", agent="Citation Verifier",
                    title=f"Crossref rate-limited — {doi} not checked",
                    detail="Crossref returned 429 (too many requests). Add CROSSREF_MAILTO in secrets for a higher rate limit.",
                    flag=False,
                ))
            else:
                findings.append(dict(
                    severity="low", agent="Citation Verifier",
                    title=f"DOI check inconclusive: {doi}",
                    detail=f"Crossref returned HTTP {r.status_code}.",
                    flag=False,
                ))
        except requests.exceptions.Timeout:
            findings.append(dict(
                severity="low", agent="Citation Verifier",
                title=f"DOI check timed out: {doi}",
                detail="Crossref did not respond within 8 seconds. Statistical and reproducibility checks are unaffected.",
                flag=False,
            ))
        except Exception as exc:
            findings.append(dict(
                severity="low", agent="Citation Verifier",
                title=f"DOI check failed: {doi}",
                detail=f"Network error: {type(exc).__name__}. Other agents continue normally.",
                flag=False,
            ))
    return findings

# ── Agent 3: Plain Language Summary ──────────────────────────────────────────
def run_pls_agent(text: str) -> str:
    if not GROQ_KEY:
        return "_No API key configured. Add GROQ\\_API\\_KEY in Streamlit secrets to enable this agent._"
    result = call_groq(
        f"Write a 3-paragraph plain-language summary of this research paper for a general audience. "
        f"No jargon. Cover: what was studied, what was found, why it matters.\n\n"
        f"Paper (first 4000 chars):\n{text[:4000]}\n\nPlain language summary:"
    )
    return result or "_Summary generation failed. Check API key._"

# ── Agent 4: Reproducibility Checker ─────────────────────────────────────────
_REPRO_CHECKS = [
    (
        "Data Availability",
        [r"data\s+avail", r"dataset\s+avail", r"zenodo", r"figshare", r"osf\.io",
         r"dryad", r"raw\s+data", r"data\s+deposit", r"data\s+are\s+avail"],
        "No data availability statement found. Readers cannot access or verify the raw data.",
    ),
    (
        "Code / Software",
        [r"code\s+avail", r"analysis\s+code", r"github\.com", r"gitlab\.com",
         r"spss", r"r\s+version\s+\d", r"python\s+\d", r"matlab", r"stata"],
        "No analysis code or software version mentioned. The analysis cannot be independently reproduced.",
    ),
    (
        "Pre-registration",
        [r"pre.?regist", r"osf\.io/\w", r"clinicaltrials\.gov", r"aspredicted",
         r"registered\s+report", r"prior\s+to\s+data\s+collect"],
        "No pre-registration found. Hypotheses may have been decided after seeing the data (HARKing).",
    ),
    (
        "Power Analysis",
        [r"power\s+anal", r"sample\s+size\s+calc", r"a\s+priori", r"g\*?power",
         r"80%?\s+power", r"\.80\s+power", r"effect\s+size.*sampl"],
        "No power analysis found. The study may be underpowered and results unreliable.",
    ),
    (
        "Materials & Procedure Detail",
        [r"stimul[ui]", r"questionnaire", r"instrument", r"appendix", r"supplementar",
         r"items?\s+were", r"scale\s+consist", r"procedure\s+was", r"measure[sd]?\s+using"],
        "Materials and procedure are not described in enough detail to replicate this study.",
    ),
]

def run_reproducibility_agent(text: str) -> tuple[int, list[dict]]:
    text_lower = text.lower()
    score = 0
    findings: list[dict] = []
    for label, patterns, fail_msg in _REPRO_CHECKS:
        passed = any(re.search(p, text_lower) for p in patterns)
        if passed:
            score += 1
            findings.append(dict(
                severity="info", agent="Reproducibility",
                title=f"{label}: Present",
                detail=f"Paper includes {label.lower()} information.",
                flag=False,
            ))
        else:
            findings.append(dict(
                severity="medium", agent="Reproducibility",
                title=f"{label}: Missing",
                detail=fail_msg,
                flag=True,
            ))
    return score, findings

# ── Agent 5: Methodology Auditor ─────────────────────────────────────────────
_STUDY_DETECT = {
    "RCT": [r"random(?:is|iz)ed\s+control", r"\bRCT\b", r"random(?:is|iz)ed\s+trial",
            r"double.blind", r"triple.blind", r"placebo.control"],
    "Meta-analysis / Systematic Review": [r"meta.anal", r"systematic\s+review", r"\bPRISMA\b",
                                           r"pool(?:ed|ing)\s+effect", r"heterogeneity"],
    "Observational": [r"\bcohort\b", r"case.control", r"cross.sectional", r"\bsurvey\b",
                      r"longitudinal\s+study", r"prospective\s+stud", r"retrospective\s+stud"],
    "Animal Study": [r"\bmice\b", r"\brats?\b", r"animal\s+(?:study|model|experiment)",
                     r"\bARRIVE\b", r"in\s+vivo\s+(?:study|experiment)"],
}
_CONSORT = [
    ("Eligibility criteria", [r"inclusion\s+crit", r"exclusion\s+crit", r"eligib\w+"]),
    ("Randomization method", [r"random(?:is|iz)ation\s+method", r"computer.generat\w+\s+random",
                               r"block\s+random", r"stratif\w+\s+random", r"random\s+number"]),
    ("Blinding / masking", [r"\bblind\w*\b", r"\bmasked?\b", r"double.blind", r"open.label"]),
    ("Primary outcome pre-specified", [r"primary\s+outcome", r"primary\s+endpoint", r"main\s+outcome"]),
    ("Sample size calculation", [r"power\s+anal", r"sample\s+size\s+calc", r"a\s+priori"]),
    ("CONSORT flow diagram", [r"\bCONSORT\b", r"flow\s+diagram", r"participant\s+flow"]),
    ("Intent-to-treat analysis", [r"intent.to.treat", r"intention.to.treat", r"\bITT\b"]),
]
_STROBE = [
    ("Study design stated", [r"cohort\s+study", r"case.control\s+study", r"cross.sectional\s+study"]),
    ("Setting and time period", [r"study\s+setting", r"data\s+collect", r"between\s+\d{4}\s+and\s+\d{4}"]),
    ("Participants described", [r"participant", r"eligib\w+\s+crit", r"recruit\w+\s+(?:from|via)"]),
    ("Confounders addressed", [r"confounder", r"covariate", r"adjust\w+\s+for", r"control\w+\s+for"]),
    ("Missing data handling", [r"missing\s+data", r"imputation", r"complete\s+case"]),
]
_PRISMA = [
    ("PRISMA / systematic review flow", [r"\bPRISMA\b", r"flow\s+diagram", r"study\s+selection"]),
    ("Database search strategy", [r"search\s+strateg", r"PubMed|Medline|EMBASE|Scopus"]),
    ("Inclusion / exclusion criteria", [r"inclusion\s+crit", r"exclusion\s+crit"]),
    ("Risk of bias assessment", [r"risk\s+of\s+bias", r"quality\s+assess", r"Cochrane", r"\bRoB\b"]),
    ("Heterogeneity assessed", [r"heterogeneity", r"I.?2\s*=", r"Q\s+statistic"]),
]
_ARRIVE = [
    ("Species / strain reported", [r"Sprague|Wistar|C57BL|BALB|nude", r"species", r"\bstrain\b"]),
    ("Ethics / IACUC approval", [r"ethics\s+(?:approv|commit)", r"\bIACUC\b", r"institutional\s+animal"]),
    ("Sample size justification", [r"power\s+anal", r"sample\s+size", r"animals?\s+per"]),
    ("Randomization described", [r"random\w+\s+assign", r"random(?:is|iz)ation"]),
    ("Blinding described", [r"blind\w*", r"masked?\b"]),
]
_GENERIC_AUDIT = [
    ("Ethics / IRB approval", [r"\bIRB\b", r"ethics\s+(?:board|commit|approv)", r"Helsinki", r"informed\s+consent"]),
    ("Conflict of interest disclosure", [r"conflict\s+of\s+interest", r"competing\s+interest", r"authors?\s+declare"]),
    ("Data availability statement", [r"data\s+avail", r"zenodo", r"figshare", r"OSF"]),
    ("Funding source disclosed", [r"fund\w+\s+(?:by|from)", r"support\w+\s+by\s+(?:a\s+)?grant"]),
]

def _detect_study_type(text: str) -> str:
    for stype, patterns in _STUDY_DETECT.items():
        if any(re.search(p, text, re.I) for p in patterns):
            return stype
    return "General"

def run_methodology_agent(text: str) -> tuple[str, str, list[dict]]:
    study_type = _detect_study_type(text)
    checklist_map = {
        "RCT": ("CONSORT", _CONSORT),
        "Meta-analysis / Systematic Review": ("PRISMA", _PRISMA),
        "Observational": ("STROBE", _STROBE),
        "Animal Study": ("ARRIVE", _ARRIVE),
        "General": ("Generic QC", _GENERIC_AUDIT),
    }
    std_name, checklist = checklist_map[study_type]
    findings = []
    for label, patterns in checklist:
        found = any(re.search(p, text, re.I) for p in patterns)
        findings.append(dict(
            severity="info" if found else "medium",
            agent="Methodology Auditor",
            title=f"{label}: {'Present' if found else 'Missing'}",
            detail=(f"Paper includes evidence of {label.lower()}." if found
                    else f"No mention of {label.lower()} found. Required by {std_name} reporting standard."),
            flag=not found,
        ))
    return study_type, std_name, findings


# ── Agent 6: Replication Predictor ────────────────────────────────────────────
_REPLICATION_FEATURES = [
    ("Large sample (n ≥ 100)",
     [r'\b[nN]\s*=\s*[1-9]\d{2,}\b', r'1[0-9]{2,}\s+(?:participant|subject|student|patient)'],
     "Small samples inflate false-positive rates; larger samples replicate more reliably."),
    ("Pre-registration present",
     [r"pre.?regist", r"osf\.io/\w", r"clinicaltrials\.gov", r"aspredicted", r"prior\s+to\s+data\s+collect"],
     "Pre-registration reduces HARKing (Hypothesizing After Results are Known)."),
    ("Effect size reported",
     [r"Cohen.?s\s+d\b", r"\bd\s*=\s*[\d.]+", r"eta.squared", r"omega.squared",
      r"Hedges.?\s*g\b", r"\bAUC\b", r"odds\s+ratio"],
     "Reporting effect sizes enables meta-analysis and replication power planning."),
    ("Multiple-comparison correction",
     [r"Bonferroni", r"\bHolm\b", r"\bFDR\b", r"false\s+discovery", r"Benjamini",
      r"adjust\w+\s+for\s+multiple", r"familywise"],
     "Without correction, multiple tests inflate the false-positive rate."),
    ("Power analysis / sample size justification",
     [r"power\s+anal", r"sample\s+size\s+calc", r"a\s+priori", r"G\*?Power", r"80%?\s+power"],
     "A priori power analysis shows the study was designed to detect the reported effect."),
    ("Open materials or code shared",
     [r"github\.com", r"gitlab\.com", r"osf\.io", r"analysis\s+code", r"open\s+(?:data|material)",
      r"zenodo", r"figshare", r"supplementar\w+\s+(?:code|data|material)"],
     "Sharing code and data allows independent verification of the analysis."),
    ("Confirmatory (not purely exploratory)",
     [r"test\w+\s+the\s+hypothes", r"confirm\w+\s+(?:our|the)\s+hypothes",
      r"(?:primary|main)\s+hypothesis\s+was", r"a\s+priori\s+hypothes"],
     "Confirmatory studies with pre-specified hypotheses replicate more often than exploratory ones."),
]

def run_replication_agent(text: str) -> tuple[int, float, list[dict]]:
    score = 0
    findings = []
    for label, patterns, rationale in _REPLICATION_FEATURES:
        found = any(re.search(p, text, re.I) for p in patterns)
        if found:
            score += 1
        findings.append(dict(
            severity="info" if found else "low",
            agent="Replication Predictor",
            title=f"{label}: {'Present' if found else 'Not detected'}",
            detail=rationale,
            flag=not found,
        ))
    # Probability estimate: 0.20 base + up to 0.60 from features (based on OSC 2015 base rate)
    prob = 0.20 + (score / len(_REPLICATION_FEATURES)) * 0.60
    return score, round(prob, 2), findings


# ── Agent 7: COI Detector ─────────────────────────────────────────────────────
_COI_NO_CONFLICT = [
    r"no\s+conflict\s+of\s+interest",
    r"no\s+competing\s+interest",
    r"authors?\s+declare\s+(?:that\s+)?(?:they\s+have\s+)?no",
    r"declare\s+no\s+(?:financial\s+)?conflict",
    r"nothing\s+to\s+disclose",
    r"no\s+financial\s+(?:ties|interest|support)",
]
_COI_INDUSTRY_SIGNALS = [
    (r"(?:funded|supported|sponsored)\s+by\s+[A-Z][a-zA-Z\s]{2,30}(?:Inc|Ltd|Corp|Pharma|Industries|LLC|GmbH)",
     "Industry-funded study"),
    (r"reports?\s+(?:grants?|funding|fees?|honoraria|personal\s+fees?)\s+(?:from|received\s+from)",
     "Author reports industry payments"),
    (r"(?:served?|acts?|work(?:s|ed)?|consult\w*)\s+(?:as\s+(?:a\s+)?)?(?:advisor|consultant|speaker)\s+for",
     "Author has advisory or speaker role"),
    (r"(?:owns?|holds?|has)\s+(?:equity|stock|shares?|patent)\s+in",
     "Author holds equity or patent"),
    (r"(?:honoraria|royalt(?:ies|y))\s+(?:received\s+)?from",
     "Author receives honoraria or royalties"),
]

def run_coi_agent(text: str) -> list[dict]:
    findings = []
    has_coi_section = bool(re.search(
        r"(?:conflict|competing)\s+of\s+interest|funding\s+(?:source|statement|acknowledgement)|author\s+contribution",
        text, re.I,
    ))
    has_no_conflict = any(re.search(p, text, re.I) for p in _COI_NO_CONFLICT)
    industry_hits = [(label, re.search(p, text, re.I).group(0)[:100])
                     for p, label in _COI_INDUSTRY_SIGNALS
                     if re.search(p, text, re.I)]

    if has_coi_section and has_no_conflict:
        findings.append(dict(
            severity="info", agent="COI Detector",
            title="No competing interests declared",
            detail="Paper includes an explicit conflict of interest statement with no conflicts disclosed.",
            flag=False,
        ))
    elif not has_coi_section:
        findings.append(dict(
            severity="medium", agent="COI Detector",
            title="No COI / funding disclosure section found",
            detail="Paper does not appear to contain a conflict of interest or funding disclosure section. "
                   "ICMJE guidelines require explicit disclosure from all authors.",
            flag=True,
        ))
    else:
        findings.append(dict(
            severity="low", agent="COI Detector",
            title="COI section present but no explicit 'no conflict' statement",
            detail="Paper has a funding or disclosure section but does not clearly state competing interests.",
            flag=True,
        ))

    for label, snippet in industry_hits:
        findings.append(dict(
            severity="medium", agent="COI Detector",
            title=f"Potential conflict: {label}",
            detail=f'Detected: "{snippet}". Human review recommended to assess materiality.',
            flag=True,
        ))

    if not findings:
        findings.append(dict(
            severity="info", agent="COI Detector",
            title="No conflict signals detected",
            detail="No industry funding or author conflict patterns were identified in the text.",
            flag=False,
        ))
    return findings


# ── Confidence Scoring ────────────────────────────────────────────────────────
def compute_confidence(all_findings: list[dict]) -> tuple[float, str]:
    """Step(38) formula: start 1.0, penalise flagged findings by severity."""
    score = 1.0
    score -= sum(0.4 for f in all_findings if f.get("flag") and f.get("severity") == "high")
    score -= sum(0.2 for f in all_findings if f.get("flag") and f.get("severity") == "medium")
    score -= sum(0.05 for f in all_findings if f.get("flag") and f.get("severity") == "low")
    score = max(0.0, score)
    label = "high" if score >= 0.7 else ("medium" if score >= 0.4 else "low")
    return round(score, 2), label


# ── PDF Export ────────────────────────────────────────────────────────────────
def export_pdf(paper_name: str, all_findings: list[dict], repro_score: int,
               repl_score: int, repl_prob: float, confidence_label: str,
               study_type: str, std_name: str, pls: str) -> bytes:
    try:
        import io as _io
        import datetime
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib.enums import TA_CENTER
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
        from reportlab.lib import colors

        buf = _io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4,
                                leftMargin=2.5*cm, rightMargin=2.5*cm,
                                topMargin=2.5*cm, bottomMargin=2.5*cm)
        styles = getSampleStyleSheet()
        title_s = ParagraphStyle("t", parent=styles["Title"], fontSize=16, spaceAfter=6, alignment=TA_CENTER)
        h2_s = ParagraphStyle("h2", parent=styles["Heading2"], fontSize=12, spaceBefore=12, spaceAfter=4)
        body_s = ParagraphStyle("b", parent=styles["Normal"], fontSize=9, leading=14, spaceAfter=4)
        small_s = ParagraphStyle("sm", parent=styles["Normal"], fontSize=8, leading=12, textColor=colors.grey)
        sev_colors = {"high": "#ef4444", "medium": "#f59e0b", "low": "#3b82f6", "info": "#9ca3af"}

        story = [
            Paragraph("PEERLESS.AI — Automated Peer Review Report", title_s),
            Paragraph(
                "<i>This report is generated by AI for informational purposes only. "
                "All findings require human expert review before any action is taken. "
                "PEERLESS.AI surfaces potential concerns — it does not make definitive judgements.</i>",
                small_s,
            ),
            Spacer(1, 0.3*cm),
            HRFlowable(width="100%", thickness=1, color=colors.HexColor("#6366f1")),
            Spacer(1, 0.2*cm),
            Paragraph(f"Paper: {paper_name}", body_s),
            Paragraph(f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M UTC')}", body_s),
            Paragraph(f"Study type detected: {study_type} ({std_name} standard)", body_s),
            Paragraph(f"Overall confidence: {confidence_label.upper()}", body_s),
            Paragraph(f"Reproducibility score: {repro_score}/5  |  Replication probability: {repl_prob:.0%}", body_s),
            Spacer(1, 0.4*cm),
        ]

        flagged = [f for f in all_findings if f.get("flag")]
        story.append(Paragraph(f"Findings Summary", h2_s))
        story.append(Paragraph(
            f"Total findings: {len(all_findings)} &nbsp;|&nbsp; "
            f"Flagged: {len(flagged)} &nbsp;|&nbsp; "
            f"High: {sum(1 for f in flagged if f['severity']=='high')} &nbsp;|&nbsp; "
            f"Medium: {sum(1 for f in flagged if f['severity']=='medium')}",
            body_s,
        ))
        story.append(Spacer(1, 0.3*cm))

        agents_order = ["Statistical Integrity", "Citation Verifier", "Reproducibility",
                        "Methodology Auditor", "Replication Predictor", "COI Detector",
                        "Plain Language Summary"]
        current_agent = None
        for f in sorted(all_findings, key=lambda x: agents_order.index(x["agent"])
                        if x["agent"] in agents_order else 99):
            if f["agent"] != current_agent:
                current_agent = f["agent"]
                story.append(Paragraph(current_agent, h2_s))
            sev = f.get("severity", "info")
            hex_c = sev_colors.get(sev, "#9ca3af")
            story.append(Paragraph(
                f'<font color="{hex_c}">[{sev.upper()}]</font> <b>{f["title"]}</b>',
                body_s,
            ))
            story.append(Paragraph(f["detail"], small_s))
            story.append(Spacer(1, 0.15*cm))

        if pls:
            story.append(Paragraph("Plain Language Summary", h2_s))
            for para in pls.split("\n\n"):
                if para.strip():
                    story.append(Paragraph(para.strip(), body_s))

        story.append(Spacer(1, 0.4*cm))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.grey))
        story.append(Paragraph(
            "DISCLAIMER: This report is AI-generated and must not be used as the sole basis for editorial "
            "decisions. All flagged concerns require verification by qualified human reviewers.",
            small_s,
        ))
        doc.build(story)
        return buf.getvalue()
    except Exception as e:
        return b""


# ── Render finding card ───────────────────────────────────────────────────────
def render_finding(f: dict):
    sev = f.get("severity", "info")
    flag_note = " &nbsp;<em>Flagged for review</em>" if f.get("flag") else ""
    st.markdown(
        f'<div class="finding-{sev}">'
        f'<span class="badge badge-{sev}">{sev.upper()}</span>&nbsp; '
        f'<strong>{f["title"]}</strong>{flag_note}<br>'
        f'<small style="color:#555;margin-top:4px;display:block">{f["detail"]}</small>'
        f'</div>',
        unsafe_allow_html=True,
    )

# ── App ───────────────────────────────────────────────────────────────────────
def main():
    st.markdown(
        '<div style="position:fixed;bottom:0;left:0;width:100%;background:#ffffff;'
        'border-top:1px solid #D6E8F4;padding:9px 28px;display:flex;align-items:center;'
        'gap:16px;z-index:9999;">'
        '<span style="font-size:0.78rem;color:#9DBFCA;">Built by '
        '<strong style="color:#1565C0;">Ahmed Irfan</strong></span>'
        '<a href="https://github.com/AhmedIrfan7/PeerLess.AI" target="_blank" '
        'style="font-size:0.78rem;color:#1565C0;text-decoration:none;font-weight:600;'
        'border:1px solid #D6E8F4;border-radius:6px;padding:2px 10px;background:#F4F8FC;">'
        'GitHub</a>'
        '</div>',
        unsafe_allow_html=True,
    )

    st.markdown("""
<div class="hero">
  <h1>PEERLESS.AI</h1>
  <p>Multi-agent scientific peer-review assistant &mdash; flags statistical errors, fake citations,
  reproducibility gaps, and methodology issues in seconds.</p>
  <span class="tag">National AI Hackathon '26</span>
  <span class="tag">Atom Camp</span>
  <span class="tag">FAST NUCES Islamabad</span>
</div>
""", unsafe_allow_html=True)

    with st.sidebar:
        st.markdown("### PEERLESS.AI")
        st.markdown("**7 Agents**")
        st.markdown("""
| Agent | |
|-------|--------|
| Statistical Integrity | on |
| Citation Verifier | on |
| Reproducibility | on |
| Methodology Auditor | on |
| Replication Predictor | on |
| COI Detector | on |
| Plain Language Summary | LLM |
""")
        st.markdown("---")
        st.markdown("**How it works**")
        st.markdown("""
1. Upload PDF or DOCX
2. 7 agents run sequentially
3. Review flagged concerns
4. All findings need human review
""")
        st.markdown("---")
        llm_ok = bool(GROQ_KEY)
        st.markdown(f"**LLM status:** {'Active' if llm_ok else 'No key — stat checks still run'}")

    uploaded = st.file_uploader("Upload a research paper (PDF or DOCX, max 20 MB)", type=["pdf", "docx"], label_visibility="visible")

    if uploaded and uploaded.size > 20 * 1024 * 1024:
        st.error("File exceeds 20 MB limit. Please upload a smaller file.")
        return

    if not uploaded:
        st.markdown('<p class="upload-hint">Upload a PDF or DOCX to run all seven agents automatically.</p>', unsafe_allow_html=True)
        st.markdown("""
<div class="agent-grid">
  <div class="agent-card">
    <div class="ac-num">1</div>
    <div class="ac-name">Statistical Integrity</div>
    <div class="ac-desc">GRIM test + p-value recomputation for t, F, and chi-squared statistics</div>
  </div>
  <div class="agent-card">
    <div class="ac-num">2</div>
    <div class="ac-name">Citation Verifier</div>
    <div class="ac-desc">Extracts every DOI and checks it live against Crossref — flags fakes and retractions</div>
  </div>
  <div class="agent-card">
    <div class="ac-num">3</div>
    <div class="ac-name">Reproducibility Checker</div>
    <div class="ac-desc">Scores 0–5 across data availability, code, pre-registration, power analysis, materials</div>
  </div>
  <div class="agent-card">
    <div class="ac-num">4</div>
    <div class="ac-name">Methodology Auditor</div>
    <div class="ac-desc">Detects study type and checks CONSORT / STROBE / PRISMA / ARRIVE compliance</div>
  </div>
  <div class="agent-card">
    <div class="ac-num">5</div>
    <div class="ac-name">Replication Predictor</div>
    <div class="ac-desc">Scores 7 evidence-based features that predict whether results will replicate</div>
  </div>
  <div class="agent-card">
    <div class="ac-num">6</div>
    <div class="ac-name">COI Detector</div>
    <div class="ac-desc">Scans for conflict of interest disclosures and industry funding signals</div>
  </div>
  <div class="agent-card">
    <div class="ac-num">7</div>
    <div class="ac-name">Plain Language Summary</div>
    <div class="ac-desc">Groq LLaMA-3.3 writes a jargon-free summary for non-specialist readers</div>
  </div>
</div>
<p style="font-size:0.78rem;color:#9DBFCA;text-align:center;margin-top:4px;">
  All findings are flagged for human review — PEERLESS.AI never makes definitive accusations.
</p>
""", unsafe_allow_html=True)
        return

    st.markdown(
        f'<div style="background:white;border:1px solid #D6E8F4;border-radius:10px;'
        f'padding:12px 16px;display:flex;align-items:center;gap:12px;margin-bottom:12px">'
        f'<span style="font-size:0.85rem;font-weight:600;color:#0C2340">{uploaded.name}</span>'
        f'<span style="font-size:0.78rem;color:#9DBFCA;margin-left:auto">{uploaded.size/1024:.1f} KB</span>'
        f'</div>',
        unsafe_allow_html=True,
    )

    if not st.button("Run Peer Review", type="primary", use_container_width=True):
        return

    pdf_bytes = uploaded.read()

    with st.spinner("Extracting text..."):
        text = extract_text(pdf_bytes, uploaded.name)

    if not text.strip():
        st.error("Could not extract text. For PDFs, make sure the file is text-based (not a scanned image). For DOCX, ensure the file is not password-protected.")
        return

    st.caption(f"{len(text):,} characters extracted")

    prog = st.progress(0, text="Running agents...")

    with st.spinner("Agent 1/7 — Statistical Integrity (GRIM + statcheck)..."):
        stat_findings = run_statistical_agent(text)
    prog.progress(14, text="Agent 2/7 — Citation Verifier...")

    with st.spinner("Agent 2/7 — Citation Verifier (Crossref)..."):
        cite_findings = run_citation_agent(text)
    prog.progress(28, text="Agent 3/7 — Reproducibility Checker...")

    with st.spinner("Agent 3/7 — Reproducibility Checker..."):
        repro_score, repro_findings = run_reproducibility_agent(text)
    prog.progress(43, text="Agent 4/7 — Methodology Auditor...")

    with st.spinner("Agent 4/7 — Methodology Auditor (CONSORT / STROBE / PRISMA)..."):
        study_type, std_name, meth_findings = run_methodology_agent(text)
    prog.progress(57, text="Agent 5/7 — Replication Predictor...")

    with st.spinner("Agent 5/7 — Replication Predictor..."):
        repl_score, repl_prob, repl_findings = run_replication_agent(text)
    prog.progress(71, text="Agent 6/7 — COI Detector...")

    with st.spinner("Agent 6/7 — COI Detector..."):
        coi_findings = run_coi_agent(text)
    prog.progress(86, text="Agent 7/7 — Plain Language Summary...")

    with st.spinner("Agent 7/7 — Plain Language Summary (Groq LLaMA)..."):
        pls = run_pls_agent(text)
    prog.progress(100, text="Done")

    all_findings = (stat_findings + cite_findings + repro_findings +
                    meth_findings + repl_findings + coi_findings)
    flagged = [f for f in all_findings if f.get("flag")]
    high = sum(1 for f in flagged if f["severity"] == "high")
    medium = sum(1 for f in flagged if f["severity"] == "medium")
    conf_score, conf_label = compute_confidence(all_findings)

    _conf_map = {"high": ("ok", "Looks Good", "No major integrity concerns detected."),
                 "medium": ("warn", "Review Recommended", "Some concerns found that need human verification."),
                 "low": ("bad", "Significant Issues", "Multiple integrity flags require careful human review.")}
    _cls, _title, _sub = _conf_map[conf_label]

    st.markdown(f"""
<div class="verdict {conf_label}">
  <div style="flex:1">
    <div class="verdict-label">Overall Confidence &mdash; {conf_label.upper()}</div>
    <div class="verdict-title">{_title}</div>
    <div class="verdict-sub">{_sub}</div>
  </div>
  <div class="verdict-stats">
    <div class="vstat"><div class="vn" style="color:#dc2626">{high}</div><div class="vl">High</div></div>
    <div class="vstat"><div class="vn" style="color:#d97706">{medium}</div><div class="vl">Medium</div></div>
    <div class="vstat"><div class="vn" style="color:#1565C0">{len(flagged)-high-medium}</div><div class="vl">Low</div></div>
    <div class="vstat"><div class="vn">{len(all_findings)}</div><div class="vl">Total</div></div>
  </div>
</div>
""", unsafe_allow_html=True)

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Flagged Issues", len(flagged))
    c2.metric("High Severity", high)
    c3.metric("Reproducibility", f"{repro_score}/5")
    c4.metric("Replication Est.", f"{repl_prob:.0%}")

    pdf_bytes = export_pdf(
        uploaded.name, all_findings, repro_score, repl_score, repl_prob,
        conf_label, study_type, std_name, pls,
    )
    if pdf_bytes:
        st.download_button(
            label="Download PDF Report",
            data=pdf_bytes,
            file_name=f"peerless_report_{uploaded.name.replace('.pdf','')}.pdf",
            mime="application/pdf",
        )

    _SEV_ORDER = ["high", "medium", "low", "info"]
    tab1, tab2, tab3, tab4, tab5, tab6, tab7 = st.tabs([
        "Statistical Integrity", "Citation Verifier", "Reproducibility",
        "Methodology Auditor", "Replication Predictor", "COI Detector",
        "Plain Language Summary",
    ])

    def _tab_header(title, count_text):
        st.markdown(
            f'<div class="tab-header"><span class="th-title">{title}</span>'
            f'<span class="th-count">{count_text}</span></div>',
            unsafe_allow_html=True,
        )

    def _score_bar(label, value, total, cls=""):
        pct = int(value / total * 100)
        st.markdown(
            f'<div class="score-row">'
            f'<span class="score-label">{label}</span>'
            f'<div class="score-track"><div class="score-fill {cls}" style="width:{pct}%"></div></div>'
            f'<span class="score-num">{value}/{total}</span>'
            f'</div>',
            unsafe_allow_html=True,
        )

    with tab1:
        flagged_stat = sum(1 for f in stat_findings if f.get("flag"))
        _tab_header("Statistical Integrity", f"{flagged_stat} flagged of {len(stat_findings)}")
        for f in sorted(stat_findings, key=lambda x: _SEV_ORDER.index(x["severity"])):
            render_finding(f)

    with tab2:
        n_dois = len(extract_dois(text))
        flagged_cite = sum(1 for f in cite_findings if f.get("flag"))
        _tab_header("Citation Verifier", f"{n_dois} DOIs checked · {flagged_cite} flagged")
        for f in sorted(cite_findings, key=lambda x: _SEV_ORDER.index(x["severity"])):
            render_finding(f)

    with tab3:
        r_cls = "ok" if repro_score >= 4 else ("warn" if repro_score >= 2 else "bad")
        r_pct_cls = r_cls
        _tab_header("Reproducibility Checker", f"{repro_score}/5 dimensions present")
        st.markdown(f"""
<div class="score-hero">
  <span class="num">{repro_score}</span><span class="denom">/5</span>
  <span class="pct {r_pct_cls}">{'Good' if repro_score>=4 else ('Partial' if repro_score>=2 else 'Poor')}</span>
</div>
<div class="score-wrap">""", unsafe_allow_html=True)
        repro_labels = ["Data Availability", "Code / Software", "Pre-registration", "Power Analysis", "Materials Detail"]
        for i, f in enumerate(repro_findings):
            lbl = repro_labels[i] if i < len(repro_labels) else f["title"]
            ok = not f.get("flag")
            _score_bar(lbl, 1 if ok else 0, 1, "ok" if ok else "bad")
        st.markdown('</div>', unsafe_allow_html=True)
        st.markdown("")
        for f in repro_findings:
            render_finding(f)

    with tab4:
        passed = sum(1 for f in meth_findings if not f.get("flag"))
        total_m = len(meth_findings)
        _tab_header(f"Methodology Auditor — {std_name}", f"{passed}/{total_m} items present")
        st.markdown(
            f'<p style="font-size:0.83rem;color:#64748b;margin:0 0 12px">'
            f'Study type detected: <strong style="color:#0C2340">{study_type}</strong>'
            f'</p>',
            unsafe_allow_html=True,
        )
        m_cls = "ok" if passed >= total_m * 0.8 else ("warn" if passed >= total_m * 0.5 else "bad")
        st.markdown('<div class="score-wrap">', unsafe_allow_html=True)
        for f in meth_findings:
            _score_bar(f["title"].replace("Present","").replace("Missing","").replace(":","").strip(),
                       0 if f.get("flag") else 1, 1,
                       "ok" if not f.get("flag") else "bad")
        st.markdown('</div>', unsafe_allow_html=True)
        st.markdown("")
        for f in meth_findings:
            render_finding(f)

    with tab5:
        repl_cls = "ok" if repl_score >= 5 else ("warn" if repl_score >= 3 else "bad")
        _tab_header("Replication Predictor", f"{repl_score}/7 features · Est. {repl_prob:.0%} probability")
        st.markdown(f"""
<div class="score-hero">
  <span class="num">{repl_score}</span><span class="denom">/7</span>
  <span class="pct {repl_cls}">{repl_prob:.0%} est.</span>
</div>
<p style="font-size:0.78rem;color:#64748b;margin:4px 0 14px">
  Estimated replication probability based on Open Science Collaboration (2015) base rates.
</p>
<div class="score-wrap">""", unsafe_allow_html=True)
        repl_labels = ["Large sample (n≥100)", "Pre-registration", "Effect size reported",
                       "Multiple-comparison correction", "Power analysis", "Open materials/code", "Confirmatory design"]
        for i, f in enumerate(repl_findings):
            lbl = repl_labels[i] if i < len(repl_labels) else f["title"]
            ok = not f.get("flag")
            _score_bar(lbl, 1 if ok else 0, 1, "ok" if ok else "bad")
        st.markdown('</div>', unsafe_allow_html=True)
        st.markdown("")
        for f in repl_findings:
            render_finding(f)

    with tab6:
        coi_flagged = sum(1 for f in coi_findings if f.get("flag"))
        _tab_header("COI Detector", f"{coi_flagged} concern(s) flagged")
        for f in sorted(coi_findings, key=lambda x: _SEV_ORDER.index(x["severity"])):
            render_finding(f)

    with tab7:
        _tab_header("Plain Language Summary", "AI-generated · Groq LLaMA-3.3")
        if pls and not pls.startswith("_"):
            st.markdown(
                f'<div style="background:white;border:1px solid #D6E8F4;border-radius:10px;'
                f'padding:20px 24px;line-height:1.7;font-size:0.92rem;color:#0C2340">{pls}</div>',
                unsafe_allow_html=True,
            )
        else:
            st.markdown(pls)

    st.markdown("""
<div class="disclaimer">
  All findings require human expert review before any action is taken.
  PEERLESS.AI surfaces potential concerns — it does not make definitive judgements.
</div>
""", unsafe_allow_html=True)



main()
